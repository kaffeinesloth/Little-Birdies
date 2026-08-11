"""
rag/extractors.py — Trích xuất text thuần từ PDF, DOCX, TXT.

Sau extract sẽ clean: bỏ nhiều newline liên tiếp, normalize whitespace.
"""
from __future__ import annotations

import logging
import re

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Nhận path của file + mime_type, trả về text thuần đã clean.

    Supported:
      - application/pdf
      - application/vnd.openxmlformats-officedocument.wordprocessingml.document (docx)
      - application/msword (doc — chỉ docx mới được, doc cần libreoffice)
      - text/plain
    """

    def extract(self, file_path: str, mime_type: str) -> str:
        mime = mime_type.lower()
        if "pdf" in mime:
            return self._extract_pdf(file_path)
        if "wordprocessingml" in mime or "msword" in mime:
            return self._extract_docx(file_path)
        if "plain" in mime or "text" in mime:
            return self._extract_txt(file_path)
        raise ValueError(f"Unsupported mime type: {mime_type}")

    # ── Extractors ──────────────────────────────────────────────────────────

    def _extract_pdf(self, path: str) -> str:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        raw = "\n\n".join(pages)
        logger.debug("PDF extracted: %d pages, %d chars", len(reader.pages), len(raw))
        return self._clean(raw)

    def _extract_docx(self, path: str) -> str:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Cũng extract text từ tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        raw = "\n\n".join(paragraphs)
        logger.debug("DOCX extracted: %d paragraphs, %d chars", len(paragraphs), len(raw))
        return self._clean(raw)

    def _extract_txt(self, path: str) -> str:
        with open(path, encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        return self._clean(raw)

    # ── Cleaning ─────────────────────────────────────────────────────────────

    @staticmethod
    def _clean(text: str) -> str:
        # Loại bỏ null bytes
        text = text.replace("\x00", "")
        # Normalize dấu xuống dòng Windows
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Bỏ nhiều newline liên tiếp (giữ tối đa 2)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Normalize whitespace ngang (tab, nhiều space)
        text = re.sub(r"[ \t]+", " ", text)
        # Bỏ space đầu/cuối mỗi dòng
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        return text.strip()
